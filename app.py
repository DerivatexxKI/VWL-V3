import streamlit as st
import openai
import os
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
from datetime import datetime
import pdfplumber
import tiktoken  # NEU: für Tokenzählung

# OpenAI-Client initialisieren
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

st.set_page_config(page_title="Volkswirtschaftliche Prognose", page_icon="📈")
st.title("📈 Volkswirtschaftliche Prognose für Regionalbanken")

# Uploadfeld für PDF- oder Word-Dateien
uploaded_files = st.file_uploader(
    "📎 Relevante PDF- oder Word-Dokumente hochladen (optional)",
    type=["pdf", "docx"],
    accept_multiple_files=True
)

# Inhalte aus hochgeladenen Dateien extrahieren
extracted_texts = []
MAX_PAGES = 25  # Maximal 25 Seiten pro PDF auslesen
if uploaded_files:
    for file in uploaded_files:
        try:
            if file.name.endswith(".pdf"):
                with pdfplumber.open(file) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages[:MAX_PAGES])
                    if not text.strip():
                        st.warning(f"⚠️ Keine lesbaren Textinhalte in PDF '{file.name}' erkannt.")
                        continue
                    extracted_texts.append(text)
                    if len(text.strip()) < 500:
                        st.warning(f"⚠️ Wenig Text erkannt in PDF '{file.name}'. Die Datei enthält möglicherweise eingescannten oder nicht maschinenlesbaren Text.")
            elif file.name.endswith(".docx"):
                from docx import Document as DocxDocument
                docx = DocxDocument(file)
                text = "\n".join([para.text for para in docx.paragraphs])
                if not text.strip():
                    st.warning(f"⚠️ Keine lesbaren Textinhalte in Word-Datei '{file.name}' erkannt.")
                    continue
                extracted_texts.append(text)
                if len(text.strip()) < 500:
                    st.warning(f"⚠️ Wenig Text erkannt in Word-Datei '{file.name}'. Bitte prüfen, ob sie inhaltlich korrekt gefüllt ist.")
        except Exception as e:
            st.error(f"❌ Fehler beim Verarbeiten der Datei '{file.name}': {e}")

# Tokenzähler definieren
encoding = tiktoken.encoding_for_model("gpt-4")
def num_tokens_from_string(string: str) -> int:
    return len(encoding.encode(string))

# Button zur Prognoseerstellung
if st.button("📈 Prognose jetzt generieren und als Word-Datei exportieren"):
    context_text = "\n\n".join(extracted_texts)

    # Ziel: maximal 115.000 Tokens für gesamten Prompt inkl. Kontext
    dummy_context = "___CONTEXT___"
    prompt_template = f"""
    Du bekommst folgende kontextuelle Dokumente als Grundlage für eine Prognose:

    {dummy_context}

    Erstelle darauf basierend eine volkswirtschaftliche Prognose für die Mittelfristplanung einer kleinen deutschen Regionalbank mit folgenden Bestandteilen:

    1. Erstelle eine Tabelle mit Planungsprämissen zu folgenden Punkten:
       - BIP-Wachstum
       - Inflation (HVPI)
       - Leitzinsen (EZB)
       - Arbeitslosenquote
       - EUR/USD Wechselkurs
       - Energiepreise (Brent, Gas)

    2. Gib für jeden dieser makroökonomischen Cluster eine **ausführlich begründete, analytisch tiefgehende** Prognose für die nächsten 3–5 Jahre ab:
       - Bruttoinlandsprodukt (BIP): Entwicklung, Sektoren, Prognoserisiken, Einflussgrößen
       - Inflation (HVPI): Kerninflation vs. Gesamt, Preisentwicklungen, Energie, Löhne
       - Arbeitsmarkt: Beschäftigungstrends, Fachkräftemangel, Sektoren
       - Geldpolitik der EZB: Begründung der Zinspolitik, Ausblick
       - Zinsstruktur (Swapkurve): Verlauf, Markterwartung, Einflussfaktoren
       - Geopolitische Risiken: Regionen, Auswirkungen, Szenarien
       - Auswirkungen auf das Bankgeschäft: Zinsüberschuss, Kreditnachfrage, Risiken
       - Wirtschaftliche Risiken: interne/externe Unsicherheiten, Stressszenarien

    Für jeden dieser Abschnitte:
    - Beginne mit einer kurzen prägnanten **Zusammenfassung** (3–4 Sätze)
    - Erstelle im Anschluss einen **detaillierten Hauptteil** mit mindestens 1800 Zeichen pro Abschnitt
    - Nutze klare Zwischenüberschriften, Aufzählungen, ggf. tabellarische Vergleiche
    - Führe Analysen, Begründungen und Beispiele an

    Verwende bevorzugt die Informationen aus den bereitgestellten Dokumenten. Wenn zu einem Bereich keine konkreten Aussagen enthalten sind, ergänze diesen Abschnitt durch fundierte eigene Einschätzungen. Zitiere bei konkreten Aussagen aus dem Upload den Wortlaut sinngemäß.

    Verwende professionelle, sachliche Sprache. Die Inhalte sollen geeignet für ein Vorstandsgremium sein.
    """
    prompt_token_count = num_tokens_from_string(prompt_template.replace(dummy_context, ""))
    st.info(f"🧠 Prompt-Grundgerüst benötigt ca. {prompt_token_count:,} Tokens")

    max_total_tokens = 115000
    available_for_context = max_total_tokens - prompt_token_count

    context_tokens = num_tokens_from_string(context_text)
    st.info(f"📄 Kontext benötigt aktuell: {context_tokens:,} Tokens (erlaubt: {available_for_context:,})")

    if context_tokens > available_for_context:
        st.warning("✂️ Kontexttext wird gekürzt, um ins Tokenlimit zu passen.")
        while num_tokens_from_string(context_text) > available_for_context:
            context_text = context_text[:-2000]  # iterativ kürzen

    prompt = prompt_template.replace(dummy_context, context_text)

    with st.spinner("Generiere Prognose..."):
        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=3000,
            )
            result = response.choices[0].message.content
            st.success("Prognose erfolgreich erstellt!")
            st.markdown(result)

            # Word-Dokument erstellen
            doc = Document()

            # Titelblatt
            doc.add_heading("Volkswirtschaftliche Mittelfristprognose", 0)
            doc.add_paragraph("Erstellt am: " + datetime.now().strftime("%d.%m.%Y"))
            doc.add_paragraph("Zielgruppe: Vorstand und Planungsteam der Regionalbank")
            doc.add_page_break()

            # Inhaltsverzeichnis-Hinweis
            doc.add_paragraph("Inhaltsverzeichnis wird automatisch generiert (in Word aktivieren).")
            doc.add_page_break()

            # Inhalt einfügen (abschnittsweise)
            for section in result.split("\n\n"):
                doc.add_paragraph(section, style='Normal')

            # Formatierung optimieren
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            # Download-Link erzeugen
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📄 Word-Bericht herunterladen",
                data=buffer,
                file_name="vwl_prognose.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Fehler beim Abruf der Prognose: {e}")
